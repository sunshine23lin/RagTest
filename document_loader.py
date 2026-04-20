# 通用文档解析器
# 支持 PDF/DOCX 格式，文本 + 表格统一输出

import os
import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document


class DocumentLoader:
    def __init__(self):
        pass

    def load_document(self, doc_path):
        """加载文档，根据文件类型自动选择解析器"""
        ext = os.path.splitext(doc_path)[1].lower()
        if ext == '.pdf':
            return self._load_pdf(doc_path)
        elif ext == '.docx':
            return self._load_docx(doc_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _load_pdf(self, pdf_path):
        """解析 PDF 文档"""
        documents = []
        all_tables = []  # 收集所有页面的表格

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # 提取表格
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table and any(row for row in table):
                        all_tables.append({
                            'page': page_num,
                            'table_idx': table_idx,
                            'data': table,
                            'source': pdf_path
                        })

                # 提取文本
                text = page.extract_text() or ""
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "page": page_num,
                            "type": "text",
                            "source": pdf_path
                        }
                    ))

        # 合并跨页表格（相同列数的连续表格）
        merged_tables = self._merge_cross_page_tables(all_tables)
        
        for table_info in merged_tables:
            table_text = self._table_to_text(table_info['data'])
            documents.append(Document(
                page_content=table_text,
                metadata={
                    "page": table_info['page'],
                    "type": "table",
                    "table_index": table_info['table_idx'],
                    "source": table_info['source']
                }
            ))

        return documents

    def _merge_cross_page_tables(self, all_tables):
        """合并跨页表格"""
        if not all_tables:
            return []

        merged = []
        current = all_tables[0]

        for i in range(1, len(all_tables)):
            next_table = all_tables[i]
            
            # 检查是否可以合并（列数相同且页面连续）
            if (len(current['data'][0]) == len(next_table['data'][0]) and 
                next_table['page'] - current['page'] <= 1):
                # 检查下一个表格的第一行是否是表头重复
                next_first_row = next_table['data'][0]
                current_first_row = current['data'][0]
                
                # 如果下一表格的第一行与当前表格第一行相似（都是表头），跳过它
                if self._rows_similar(current_first_row, next_first_row):
                    # 只合并数据行，跳过表头
                    current['data'].extend(next_table['data'][1:])
                else:
                    # 不相似，直接合并所有行
                    current['data'].extend(next_table['data'])
                
                current['page'] = max(current['page'], next_table['page'])
            else:
                merged.append(current)
                current = next_table

        merged.append(current)
        return merged

    def _rows_similar(self, row1, row2):
        """检查两行是否相似（都是表头）"""
        if len(row1) != len(row2):
            return False
        
        # 计算相同位置非空单元格的重合度
        matches = 0
        total = 0
        for c1, c2 in zip(row1, row2):
            if c1 or c2:  # 至少一个非空
                total += 1
                if c1 and c2 and (c1.strip() == c2.strip() or 
                                  (not c1.strip()) or (not c2.strip())):
                    matches += 1
        
        return total > 0 and matches / total > 0.5

    def _load_docx(self, docx_path):
        """解析 DOCX 文档"""
        documents = []
        doc = DocxDocument(docx_path)

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "paragraph_index": para_idx,
                        "type": "text",
                        "source": docx_path
                    }
                ))

        for table_idx, table in enumerate(doc.tables):
            table_text = self._docx_table_to_text(table)
            documents.append(Document(
                page_content=table_text,
                metadata={
                    "table_index": table_idx,
                    "type": "table",
                    "source": docx_path
                }
            ))

        return documents

    def _table_to_text(self, table):
        """将 PDF 表格转换为 Markdown 格式文本，支持合并表头"""
        if not table or len(table) == 0:
            return ""

        rows = []
        
        # 清理单元格内的换行符
        def clean_cell(cell):
            if cell:
                return cell.strip().replace('\n', '').replace('\r', '')
            return ""
        
        # 检测是否为合并表头（第一行有空单元格，第二行有子列标签）
        has_merged_header = False
        if len(table) >= 2:
            first_row = table[0]
            second_row = table[1]
            # 如果第一行存在空单元格且第二行对应位置有值，说明是合并表头
            empty_in_first = sum(1 for c in first_row if not c or not c.strip())
            filled_in_second = sum(1 for c in second_row if c and c.strip())
            if empty_in_first > 0 and filled_in_second > 0:
                has_merged_header = True

        if has_merged_header:
            # 合并两行表头：将第一行的组名填充到第二行的空单元格中
            first_row = table[0]
            second_row = table[1]
            merged_header = []
            current_group = ""
            for i in range(len(first_row)):
                cell1 = clean_cell(first_row[i])
                cell2 = clean_cell(second_row[i])
                if cell1:
                    current_group = cell1
                if cell2:
                    merged_header.append(f"{current_group}_{cell2}" if current_group and current_group != cell2 else cell2)
                else:
                    merged_header.append(current_group)
            
            header_text = "| " + " | ".join(merged_header) + " |"
            rows.append(header_text)
            rows.append("| " + " | ".join("---" for _ in merged_header) + " |")
            
            # 数据行从第三行开始
            for row in table[2:]:
                row_text = "| " + " | ".join(clean_cell(cell) for cell in row) + " |"
                rows.append(row_text)
        else:
            # 普通表头
            header = table[0]
            header_text = "| " + " | ".join(clean_cell(cell) for cell in header) + " |"
            rows.append(header_text)
            rows.append("| " + " | ".join("---" for _ in header) + " |")
            for row in table[1:]:
                row_text = "| " + " | ".join(clean_cell(cell) for cell in row) + " |"
                rows.append(row_text)

        return "\n".join(rows)

    def _docx_table_to_text(self, table):
        """将 DOCX 表格转换为 Markdown 格式文本"""
        if not table or len(table.rows) == 0:
            return ""

        rows = []
        # 表头
        header = table.rows[0]
        header_text = "| " + " | ".join(cell.text for cell in header.cells) + " |"
        rows.append(header_text)
        rows.append("| " + " | ".join("---" for _ in header.cells) + " |")

        # 数据行
        for row in table.rows[1:]:
            row_text = "| " + " | ".join(cell.text for cell in row.cells) + " |"
            rows.append(row_text)

        return "\n".join(rows)
